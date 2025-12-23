# grammar_checker/exporters/docx_exporter.py
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.http import HttpResponse
from datetime import datetime
import io

def _set_run_color(run, rgb_color):
    run.font.color.rgb = RGBColor(*rgb_color)

def export_to_docx(request_id: int, original_text: str, corrected_text: str):
    doc = Document()

    # Tiêu đề
    # title = doc.add_paragraph("GRAMMAR CHECKER – KẾT QUẢ SỬA LỖI")
    # title.runs[0].font.size = Pt(20)
    # title.runs[0].bold = True
    # _set_run_color(title.runs[0], (25, 135, 84))  # Xanh lá
    # title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #
    # doc.add_paragraph(f"Mã yêu cầu: #{request_id}")
    # doc.add_paragraph(f"Thời gian: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    # doc.add_page_break()

    # Văn bản gốc
    # doc.add_heading('Văn bản gốc', level=2)
    # p1 = doc.add_paragraph(original_text)
    # for run in p1.runs:
    #     run.font.name = 'Times New Roman'
    #     run.font.size = Pt(12)

    # doc.add_page_break()

    # Văn bản đã sửa
    # doc.add_heading('Văn bản đã được chỉnh sửa', level=2)
    p2 = doc.add_paragraph(corrected_text)
    for run in p2.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        # _set_run_color(run, (0, 120, 0))  # Xanh đậm đẹp

    # doc.add_paragraph('\n\n© 2025 Grammar Checker Pro – Đồ án xuất sắc nhất khoa', style='Intense Quote')

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="Grammar_Check_{request_id}.docx"'
    return response